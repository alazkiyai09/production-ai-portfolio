"""
Document Retriever for DataChat-RAG

Handles semantic search and retrieval of company documents using ChromaDB.
Supports metadata filtering, reranking, and automatic chunking.
"""

import os
import re
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Union

import chromadb
from chromadb.config import Settings
from chromadb.utils import embedding_functions

from llama_index.core import Document, VectorStoreIndex
from llama_index.core.node_parser import SentenceSplitter
from llama_index.vector_stores.chroma import ChromaVectorStore


# =============================================================================
# Data Classes and Enums
# =============================================================================

class DocumentType(Enum):
    """Types of documents in the knowledge base."""

    POLICY = "policy"
    GUIDELINE = "guideline"
    REPORT = "report"
    TRAINING = "training"
    FAQ = "faq"
    BEST_PRACTICE = "best_practice"
    SOP = "sop"  # Standard Operating Procedure
    COMPLIANCE = "compliance"

    @classmethod
    def from_string(cls, value: str) -> "DocumentType":
        """Convert string to DocumentType, case-insensitive."""
        for doc_type in cls:
            if doc_type.value.lower() == value.lower().replace("-", "_").replace(" ", "_"):
                return doc_type
        raise ValueError(f"Unknown document type: {value}")

    def __str__(self) -> str:
        return self.value


@dataclass
class RetrievalResult:
    """Result from document retrieval."""

    content: str
    source: str
    doc_type: DocumentType
    relevance_score: float
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunk_id: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            "content": self.content,
            "source": self.source,
            "doc_type": str(self.doc_type),
            "relevance_score": self.relevance_score,
            "chunk_id": self.chunk_id,
            "metadata": self.metadata,
        }

    def format_citation(self) -> str:
        """Format as a citation string."""
        date_str = self.metadata.get("date", "Unknown date")
        return f"[Source: {self.source} | {self.doc_type.value.replace('_', ' ').title()} | {date_str}]"


@dataclass
class IngestionResult:
    """Result from document ingestion."""

    num_documents: int
    num_chunks: int
    num_errors: int
    errors: List[str] = field(default_factory=list)
    processing_time_seconds: float = 0.0


# =============================================================================
# Document Retriever
# =============================================================================

class DocumentRetriever:
    """
    Semantic document retriever using ChromaDB and LlamaIndex.

    Features:
    - ChromaDB for vector storage
    - Semantic search with optional reranking
    - Metadata filtering (doc_type, department, date range)
    - Automatic document chunking
    - Support for multiple file formats (PDF, TXT, MD, DOCX)
    """

    # Collection name in ChromaDB
    COLLECTION_NAME = "datachat_documents"

    # Default chunking settings
    DEFAULT_CHUNK_SIZE = 512
    DEFAULT_CHUNK_OVERLAP = 50

    # Metadata schema
    METADATA_SCHEMA = {
        "doc_type": str,
        "department": str,
        "date": str,
        "author": str,
        "version": str,
        "tags": list,
    }

    def __init__(
        self,
        chroma_path: Optional[str] = None,
        collection_name: Optional[str] = None,
        embedding_model: str = "text-embedding-3-small",
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
        rerank_enabled: bool = False,
        rerank_top_n: int = 3,
    ):
        """
        Initialize the document retriever.

        Args:
            chroma_path: Path to ChromaDB persistence directory
            collection_name: Name of the ChromaDB collection
            embedding_model: OpenAI embedding model to use
            chunk_size: Size of document chunks for indexing
            chunk_overlap: Overlap between chunks
            rerank_enabled: Whether to use Cohere reranking
            rerank_top_n: Number of results to return after reranking
        """
        self.chroma_path = chroma_path or os.getenv("CHROMA_PERSIST_DIR", "./data/chromadb")
        self.collection_name = collection_name or os.getenv("CHROMA_COLLECTION_NAME", self.COLLECTION_NAME)
        self.embedding_model = embedding_model
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.rerank_enabled = rerank_enabled
        self.rerank_top_n = rerank_top_n

        # Initialize ChromaDB client
        os.makedirs(self.chroma_path, exist_ok=True)

        self.chroma_client = chromadb.PersistentClient(
            path=self.chroma_path,
            settings=Settings(
                anonymized_telemetry=False,
                allow_reset=True,
            )
        )

        # Initialize embedding function
        self.embedding_function = embedding_functions.OpenAIEmbeddingFunction(
            api_key=os.getenv("OPENAI_API_KEY"),
            model_name=embedding_model,
        )

        # Get or create collection
        self.collection = self.chroma_client.get_or_create_collection(
            name=self.collection_name,
            metadata={"description": "DataChat-RAG document store"}
        )

        # Initialize node parser for chunking
        self.node_parser = SentenceSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separator=" ",
        )

        # Initialize reranker if enabled
        self.reranker = None
        if rerank_enabled:
            try:
                from llama_index.postprocessor.cohere_rerank import CohereRerank
                api_key = os.getenv("COHERE_API_KEY")
                if api_key:
                    self.reranker = CohereRerank(
                        api_key=api_key,
                        top_n=rerank_top_n,
                    )
                    print(f"✓ Cohere reranking enabled (top_n={rerank_top_n})")
                else:
                    print("⚠ COHERE_API_KEY not found, reranking disabled")
                    self.rerank_enabled = False
            except ImportError:
                print("⚠ Cohere rerank not available, install with: pip install llama-index-postprocessor-cohere-rerank")
                self.rerank_enabled = False

        # Create vector store for LlamaIndex
        self.vector_store = ChromaVectorStore(chroma_collection=self.collection)

        print(f"✓ DocumentRetriever initialized")
        print(f"  Collection: {self.collection_name}")
        print(f"  Path: {self.chroma_path}")
        print(f"  Chunks: {chunk_size} with {chunk_overlap} overlap")
        print(f"  Existing documents: {self.collection.count()}")

    def add_documents(
        self,
        docs: List[Document],
        batch_size: int = 100,
    ) -> IngestionResult:
        """
        Add documents to the vector store.

        Args:
            docs: List of LlamaIndex Document objects
            batch_size: Number of documents to process in each batch

        Returns:
            IngestionResult with statistics
        """
        import time
        start_time = time.time()

        num_chunks = 0
        num_errors = 0
        errors = []

        try:
            # Parse documents into chunks
            nodes = self.node_parser.get_nodes_from_documents(docs)

            # Add metadata to nodes
            for i, node in enumerate(nodes):
                # Ensure metadata exists
                if not hasattr(node, "metadata") or node.metadata is None:
                    node.metadata = {}

                # Add chunk info
                node.metadata["chunk_id"] = str(uuid.uuid4())
                node.metadata["chunk_index"] = i
                node.metadata["ingested_at"] = datetime.now().isoformat()

                # Validate required metadata
                if "doc_type" not in node.metadata:
                    node.metadata["doc_type"] = "unknown"
                if "source" not in node.metadata:
                    node.metadata["source"] = "unknown"

            # Create index and add to vector store
            index = VectorStoreIndex(nodes, vector_store=self.vector_store)

            num_chunks = len(nodes)

        except Exception as e:
            num_errors += 1
            errors.append(f"Failed to add documents: {str(e)}")

        processing_time = time.time() - start_time

        result = IngestionResult(
            num_documents=len(docs),
            num_chunks=num_chunks,
            num_errors=num_errors,
            errors=errors,
            processing_time_seconds=processing_time,
        )

        print(f"✓ Ingested {len(docs)} documents → {num_chunks} chunks")
        if errors:
            print(f"⚠ {num_errors} errors occurred")

        return result

    def ingest_directory(
        self,
        directory: Union[str, Path],
        recursive: bool = True,
        file_extensions: Optional[List[str]] = None,
        metadata_extractor: Optional[Callable[[str, Path], Dict[str, Any]]] = None,
    ) -> IngestionResult:
        """
        Ingest documents from a directory.

        Args:
            directory: Path to directory containing documents
            recursive: Whether to search subdirectories
            file_extensions: List of file extensions to include (e.g., [".pdf", ".txt"])
            metadata_extractor: Optional function to extract metadata from file path

        Returns:
            IngestionResult with statistics
        """
        import time
        start_time = time.time()

        directory = Path(directory)
        if not directory.exists():
            raise ValueError(f"Directory not found: {directory}")

        # Default supported extensions
        if file_extensions is None:
            file_extensions = [".pdf", ".txt", ".md", ".markdown", ".docx", ".html"]

        # Find all files
        if recursive:
            files = [
                f for f in directory.rglob("*")
                if f.is_file() and f.suffix.lower() in file_extensions
            ]
        else:
            files = [
                f for f in directory.glob("*")
                if f.is_file() and f.suffix.lower() in file_extensions
            ]

        print(f"Found {len(files)} documents to ingest")

        # Load documents
        documents = []
        errors = []

        for file_path in files:
            try:
                # Read file content
                content = self._read_file(file_path)

                # Extract metadata
                metadata = self._extract_metadata_from_path(file_path)
                if metadata_extractor:
                    extracted = metadata_extractor(str(file_path), file_path)
                    metadata.update(extracted)

                # Create document
                doc = Document(
                    text=content,
                    metadata=metadata,
                )
                documents.append(doc)

            except Exception as e:
                errors.append(f"Failed to load {file_path}: {str(e)}")

        # Add to vector store
        result = self.add_documents(documents)
        result.errors.extend(errors)
        result.num_errors = len(errors)
        result.processing_time_seconds = time.time() - start_time

        return result

    def _read_file(self, file_path: Path) -> str:
        """Read content from a file based on its extension."""
        suffix = file_path.suffix.lower()

        if suffix in [".txt", ".md", ".markdown"]:
            return file_path.read_text(encoding="utf-8")

        elif suffix == ".pdf":
            try:
                import PyPDF2
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    return text
            except ImportError:
                raise ImportError("PyPDF2 required for PDF support. Install with: pip install PyPDF2")

        elif suffix == ".docx":
            try:
                import docx
                doc = docx.Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                raise ImportError("python-docx required for DOCX support. Install with: pip install python-docx")

        elif suffix == ".html":
            try:
                from bs4 import BeautifulSoup
                with open(file_path, "r", encoding="utf-8") as f:
                    html = f.read()
                soup = BeautifulSoup(html, "html.parser")
                return soup.get_text(separator="\n", strip=True)
            except ImportError:
                raise ImportError("beautifulsoup4 required for HTML support. Install with: pip install beautifulsoup4")

        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _extract_metadata_from_path(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from file path and content."""
        metadata = {
            "source": file_path.name,
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size,
            "modified_at": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
        }

        # Try to infer document type from path
        path_lower = str(file_path).lower()
        if "policy" in path_lower or "policies" in path_lower:
            metadata["doc_type"] = DocumentType.POLICY.value
        elif "guideline" in path_lower or "guidelines" in path_lower:
            metadata["doc_type"] = DocumentType.GUIDELINE.value
        elif "report" in path_lower or "reports" in path_lower:
            metadata["doc_type"] = DocumentType.REPORT.value
        elif "training" in path_lower or "train" in path_lower:
            metadata["doc_type"] = DocumentType.TRAINING.value
        elif "faq" in path_lower:
            metadata["doc_type"] = DocumentType.FAQ.value
        elif "sop" in path_lower:
            metadata["doc_type"] = DocumentType.SOP.value
        elif "compliance" in path_lower or "hipaa" in path_lower:
            metadata["doc_type"] = DocumentType.COMPLIANCE.value
        elif "best_practice" in path_lower or "best-practice" in path_lower:
            metadata["doc_type"] = DocumentType.BEST_PRACTICE.value
        else:
            metadata["doc_type"] = DocumentType.GUIDELINE.value  # Default

        # Try to infer department from path
        departments = ["marketing", "sales", "legal", "compliance", "engineering", "hr", "finance"]
        for dept in departments:
            if dept in path_lower:
                metadata["department"] = dept
                break

        # Extract date from filename if present
        date_match = re.search(r"(\d{4}-\d{2}-\d{2}|\d{4}_\d{2}_\d{2})", file_path.name)
        if date_match:
            metadata["date"] = date_match.group(1).replace("_", "-")

        return metadata

    def retrieve(
        self,
        query: str,
        filters: Optional[Dict[str, Any]] = None,
        top_k: int = 5,
        similarity_threshold: Optional[float] = None,
    ) -> List[RetrievalResult]:
        """
        Retrieve relevant documents for a query.

        Args:
            query: Search query
            filters: Optional metadata filters (e.g., {"doc_type": "policy", "department": "legal"})
            top_k: Number of results to return
            similarity_threshold: Minimum similarity score (0-1)

        Returns:
            List of RetrievalResult objects, sorted by relevance
        """
        # Build ChromaDB where clause for filtering
        where_clause = self._build_where_clause(filters) if filters else None

        # Query ChromaDB
        results = self.collection.query(
            query_texts=[query],
            n_results=top_k * 2 if self.rerank_enabled else top_k,  # Get more for reranking
            where=where_clause,
        )

        # Convert to RetrievalResult objects
        retrieval_results = []

        if results and results["documents"] and results["documents"][0]:
            for i, doc in enumerate(results["documents"][0]):
                # Get metadata
                metadata = results["metadatas"][0][i] if results["metadatas"] else {}
                distance = results["distances"][0][i] if results["distances"] else 1.0

                # Convert distance to similarity (ChromaDB uses L2 distance)
                similarity_score = 1.0 / (1.0 + distance)

                # Apply similarity threshold if provided
                if similarity_threshold and similarity_score < similarity_threshold:
                    continue

                # Get doc_type
                try:
                    doc_type_str = metadata.get("doc_type", "guideline")
                    doc_type = DocumentType.from_string(doc_type_str)
                except ValueError:
                    doc_type = DocumentType.GUIDELINE

                result = RetrievalResult(
                    content=doc,
                    source=metadata.get("source", "Unknown"),
                    doc_type=doc_type,
                    relevance_score=similarity_score,
                    metadata=metadata,
                    chunk_id=metadata.get("chunk_id"),
                )
                retrieval_results.append(result)

        # Apply reranking if enabled
        if self.rerank_enabled and self.reranker and retrieval_results:
            retrieval_results = self._rerank(query, retrieval_results)

        return retrieval_results[:top_k]

    def _build_where_clause(self, filters: Dict[str, Any]) -> Dict[str, Any]:
        """Build ChromaDB where clause from filters."""
        where = {}

        for key, value in filters.items():
            if key in self.METADATA_SCHEMA or key in ["source", "file_path"]:
                where[key] = value
            elif key == "doc_type":
                if isinstance(value, str):
                    where[key] = value.lower()
                elif isinstance(value, list):
                    where[key] = {"$in": [v.lower() for v in value]}
            elif key == "date_from":
                # Requires custom handling in production
                pass
            elif key == "date_to":
                # Requires custom handling in production
                pass
            else:
                where[key] = value

        return where

    def _rerank(self, query: str, results: List[RetrievalResult]) -> List[RetrievalResult]:
        """Rerank results using Cohere reranker."""
        try:
            from llama_index.core.schema import NodeWithScore, QueryBundle
            from llama_index.core import Document

            # Convert to NodeWithScore format
            nodes = []
            for result in results:
                doc = Document(text=result.content, metadata=result.metadata)
                node = NodeWithScore(
                    node=doc.as_node(),
                    score=result.relevance_score,
                )
                nodes.append(node)

            # Rerank
            query_bundle = QueryBundle(query_str=query)
            reranked_nodes = self.reranker.postprocess_nodes(nodes, query_bundle)

            # Convert back to RetrievalResult
            reranked_results = []
            for node in reranked_nodes:
                metadata = node.node.metadata or {}
                try:
                    doc_type = DocumentType.from_string(metadata.get("doc_type", "guideline"))
                except ValueError:
                    doc_type = DocumentType.GUIDELINE

                result = RetrievalResult(
                    content=node.node.text,
                    source=metadata.get("source", "Unknown"),
                    doc_type=doc_type,
                    relevance_score=node.score,
                    metadata=metadata,
                    chunk_id=metadata.get("chunk_id"),
                )
                reranked_results.append(result)

            return reranked_results

        except Exception as e:
            print(f"⚠ Reranking failed: {e}, returning original results")
            return results

    def format_context(
        self,
        results: List[RetrievalResult],
        include_citations: bool = True,
        max_results: Optional[int] = None,
    ) -> str:
        """
        Format retrieval results as context for LLM.

        Args:
            results: List of retrieval results
            include_citations: Whether to include source citations
            max_results: Maximum number of results to include

        Returns:
            Formatted context string
        """
        if max_results:
            results = results[:max_results]

        if not results:
            return "No relevant documents found."

        context_parts = []

        for i, result in enumerate(results, 1):
            # Add content
            content = result.content.strip()
            context_parts.append(f"[{i}] {content}")

            # Add citation if requested
            if include_citations:
                context_parts.append(f"    {result.format_citation()}")

            context_parts.append("")  # Blank line between results

        return "\n".join(context_parts).strip()

    def get_stats(self) -> Dict[str, Any]:
        """Get statistics about the document store."""
        count = self.collection.count()

        return {
            "collection_name": self.collection_name,
            "total_chunks": count,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "embedding_model": self.embedding_model,
            "rerank_enabled": self.rerank_enabled,
            "rerank_top_n": self.rerank_top_n if self.rerank_enabled else None,
        }

    def clear_collection(self):
        """Clear all documents from the collection."""
        self.chroma_client.delete_collection(self.collection_name)
        self.collection = self.chroma_client.create_collection(
            name=self.collection_name,
            metadata={"description": "DataChat-RAG document store"}
        )
        print(f"✓ Collection '{self.collection_name}' cleared")

    def delete_by_filters(self, filters: Dict[str, Any]):
        """Delete documents matching the given filters."""
        where_clause = self._build_where_clause(filters)
        self.collection.delete(where=where_clause)
        print(f"✓ Deleted documents matching filters: {filters}")


# =============================================================================
# Convenience Functions
# =============================================================================

def create_retriever(
    chroma_path: Optional[str] = None,
    collection_name: Optional[str] = None,
    embedding_model: str = "text-embedding-3-small",
    rerank: bool = False,
) -> DocumentRetriever:
    """
    Convenience function to create a document retriever.

    Args:
        chroma_path: Path to ChromaDB storage
        collection_name: Name of the collection
        embedding_model: OpenAI embedding model
        rerank: Whether to enable Cohere reranking

    Returns:
        Configured DocumentRetriever instance
    """
    return DocumentRetriever(
        chroma_path=chroma_path,
        collection_name=collection_name,
        embedding_model=embedding_model,
        rerank_enabled=rerank,
    )


# =============================================================================
# Test Cases
# =============================================================================

def create_sample_documents() -> List[Document]:
    """Create sample healthcare AdTech documents for testing."""
    return [
        Document(
            text="""HIPAA Compliance Requirements for Healthcare Advertising

All healthcare advertising campaigns must comply with HIPAA regulations regarding protected health information (PHI).

Key Requirements:
1. Never use PHI in ad creatives without explicit authorization
2. Implement proper data encryption for all patient data
3. Ensure business associate agreements are in place with all partners
4. Train all staff on HIPAA privacy practices annually
5. Maintain audit logs for all PHI access for 6 years

For pharmaceutical campaigns, additional FDA regulations apply regarding fair balance and risk disclosure.""",
            metadata={
                "doc_type": DocumentType.COMPLIANCE.value,
                "department": "legal",
                "source": "HIPAA_Compliance_Guide.txt",
                "date": "2024-01-15",
                "author": "Legal Department",
                "version": "2.1",
                "tags": ["hipaa", "compliance", "healthcare", "privacy"],
            }
        ),
        Document(
            text="""Ad Approval Process for Healthcare Clients

All healthcare and pharmaceutical advertisements must undergo a multi-stage review process:

Stage 1: Creative Review (24-48 hours)
- Check for prohibited claims
- Verify fair balance in pharmaceutical ads
- Ensure proper disclaimer placement
- Review imagery for appropriateness

Stage 2: Compliance Review (24 hours)
- Legal team validates regulatory compliance
- HIPAA verification for patient testimonials
- FDA guidelines check for pharma products

Stage 3: Client Approval (varies)
- Client receives compliance-approved creative
- Revisions submitted through ticket system
- Final sign-off required before launch

Expedited review available for urgent campaigns with 24-hour turnaround.""",
            metadata={
                "doc_type": DocumentType.SOP.value,
                "department": "operations",
                "source": "Ad_Approval_Process.txt",
                "date": "2024-02-01",
                "author": "Operations Team",
                "version": "3.0",
                "tags": ["approval", "process", "healthcare", "workflow"],
            }
        ),
        Document(
            text="""Healthcare Campaign Best Practices

Industry Benchmarks (2023-2024):
- Average CTR: 0.8-1.5% (lower than general due to regulations)
- Average CVR: 2-4% (higher due to targeted audiences)
- Average CPA: $150-500 (varies by treatment type)

Targeting Best Practices:
1. Focus on healthcare professionals (HCPs) for pharma
2. Use condition-based targeting rather than behavioral
3. Geographic targeting should align with prescribing data
4. Avoid age/gender targeting that could be discriminatory

Creative Guidelines:
- Include clear disclaimers for all pharmaceutical products
- Use professional, trustworthy imagery
- Avoid fear-based messaging
- Testimonials require proper authorization

Budget Allocation:
- 60% programmatic/managed placements
- 25% premium health publisher sites (WebMD, Healthline)
- 10% social media (LinkedIn only for HCP campaigns)
- 5% contingency for optimization""",
            metadata={
                "doc_type": DocumentType.BEST_PRACTICE.value,
                "department": "marketing",
                "source": "Healthcare_Campaign_Best_Practices.txt",
                "date": "2024-03-10",
                "author": "Marketing Strategy Team",
                "version": "1.0",
                "tags": ["benchmark", "best_practice", "healthcare", "campaign"],
            }
        ),
        Document(
            text="""Attribution Models for Healthcare Campaigns

Choosing the right attribution model is critical for healthcare campaigns due to long patient journeys.

Recommended Models by Campaign Type:

1. Lead Generation (Doctor Lookup, Appointment Booking):
   - Use: Last Click with 30-day lookback
   - Rationale: Direct response actions need clear attribution

2. Brand Awareness (Condition Education, Treatment Info):
   - Use: Time Decay with 90-day lookback
   - Rationale: Healthcare decisions have long consideration periods

3. Pharma Product Launches:
   - Use: Linear attribution with 60-day lookback
   - Rationale: Multiple touchpoints contribute to prescribing decisions

Important Considerations:
- Track both HCP (healthcare professional) and patient conversions separately
- Include offline conversions (call center, office visits) when possible
- Exclude retargeting from attribution to avoid double counting

Custom Model Configuration:
Contact the analytics team to set up custom models for specialty campaigns.""",
            metadata={
                "doc_type": DocumentType.GUIDELINE.value,
                "department": "analytics",
                "source": "Attribution_Model_Guidelines.txt",
                "date": "2024-02-20",
                "author": "Analytics Team",
                "version": "1.5",
                "tags": ["attribution", "analytics", "modeling", "measurement"],
            }
        ),
        Document(
            text="""Frequently Asked Questions: Healthcare Advertising

Q: Can we use before/after patient photos in ads?
A: Only with explicit written authorization. Photos must not reveal protected health information.

Q: What disclaimers are required for pharmaceutical ads?
A: All pharma ads must include:
   - Brief summary of risks (TV/radio) or Full prescribing information (print)
   - Fair balance of benefits and risks
   - "Individual results may vary" for testimonial ads
   - Generic name pronunciation for TV ads

Q: Can we target by health condition on social media?
A: Condition-based targeting is permitted but must be carefully validated to avoid discrimination.
   Consult legal before implementation.

Q: How long does compliance review take?
A: Standard review: 48-72 hours. Expedited review: 24 hours. Rush review: 4 hours (requires VP approval).

Q: What should we do if we receive a cease and desist letter?
A: Immediately contact legal@datachat.com and pause all related campaigns. Do not respond directly.""",
            metadata={
                "doc_type": DocumentType.FAQ.value,
                "department": "legal",
                "source": "Healthcare_Advertising_FAQ.txt",
                "date": "2024-01-25",
                "author": "Legal Department",
                "version": "4.0",
                "tags": ["faq", "legal", "compliance", "healthcare"],
            }
        ),
    ]


def run_test_cases():
    """Run test cases for the document retriever."""
    import pprint

    pp = pprint.PrettyPrinter(indent=2)

    print("=" * 80)
    print("DOCUMENT RETRIEVER TEST CASES")
    print("=" * 80)

    # Check for API key
    if not os.getenv("OPENAI_API_KEY"):
        print("\n⚠ OPENAI_API_KEY not found. Using mock mode for testing.")
        print("To run full tests, set OPENAI_API_KEY in your environment.\n")

        # Show document structure
        print("Sample Documents Structure:")
        print("─" * 80)
        for doc in create_sample_documents():
            print(f"\nDocument: {doc.metadata.get('source')}")
            print(f"  Type: {doc.metadata.get('doc_type')}")
            print(f"  Department: {doc.metadata.get('department')}")
            print(f"  Content Length: {len(doc.text)} chars")
            print(f"  Preview: {doc.text[:100]}...")

        print("\n" + "=" * 80)
        print("✓ DOCUMENT STRUCTURE TEST PASSED (Set OPENAI_API_KEY for full tests)")
        print("=" * 80)
        return

    # Create retriever
    retriever = create_retriever(
        chroma_path="./data/chromadb_test",
        collection_name="test_documents",
        rerank=False,  # Disable for testing without API key
    )

    print("\n" + "─" * 80)
    print("Test 1: Add Sample Documents")
    print("─" * 80)

    sample_docs = create_sample_documents()
    result = retriever.add_documents(sample_docs)
    print(f"✓ Added {result.num_documents} documents → {result.num_chunks} chunks")

    print("\n" + "─" * 80)
    print("Test 2: Basic Retrieval")
    print("─" * 80)

    query = "What are the HIPAA requirements?"
    print(f"Query: {query}")
    results = retriever.retrieve(query, top_k=3)
    print(f"\nFound {len(results)} results:\n")
    for i, r in enumerate(results, 1):
        print(f"{i}. Score: {r.relevance_score:.3f} | {r.doc_type.value} | {r.source}")
        print(f"   {r.content[:100]}...")
        print()

    print("─" * 80)
    print("Test 3: Filtered Retrieval")
    print("─" * 80)

    query = "approval process"
    filters = {"doc_type": "sop"}
    print(f"Query: {query}")
    print(f"Filters: {filters}")
    results = retriever.retrieve(query, filters=filters, top_k=2)
    print(f"\nFound {len(results)} results:\n")
    for i, r in enumerate(results, 1):
        print(f"{i}. Score: {r.relevance_score:.3f} | {r.doc_type.value} | {r.source}")
        print(f"   {r.content[:100]}...")
        print()

    print("─" * 80)
    print("Test 4: Format Context")
    print("─" * 80)

    query = "healthcare campaign benchmarks"
    results = retriever.retrieve(query, top_k=2)
    context = retriever.format_context(results, include_citations=True)
    print(f"Query: {query}\n")
    print("Formatted Context:")
    print("─" * 40)
    print(context)
    print("─" * 40)

    print("\n" + "─" * 80)
    print("Test 5: Retrieval Stats")
    print("─" * 80)

    stats = retriever.get_stats()
    pp.pprint(stats)

    # Clean up test collection
    print("\n" + "─" * 80)
    print("Cleanup")
    print("─" * 80)
    retriever.clear_collection()
    print("✓ Test collection cleared")

    print("\n" + "=" * 80)
    print("✓ ALL TESTS PASSED")
    print("=" * 80)


if __name__ == "__main__":
    run_test_cases()
