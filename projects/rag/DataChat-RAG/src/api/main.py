"""
FastAPI Application for DataChat-RAG

REST API for the healthcare AdTech RAG system.
"""

import asyncio
import json
import logging
import os
import sys
import time
import uuid
from contextlib import asynccontextmanager
from datetime import datetime
from typing import Any, Dict, List, Optional, AsyncGenerator

from fastapi import (
    FastAPI,
    HTTPException,
    UploadFile,
    File,
    Form,
    BackgroundTasks,
    status,
    Request,
    Depends,
    Body,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel, Field
import uvicorn

# Shared modules
from shared.security import SensitiveDataFilter, install_security_filter
from shared.rate_limit import limiter, rate_limit_exception_handler, RateLimitExceeded
from shared.auth import (
    get_current_user,
    require_role,
    require_admin,
    create_access_token,
    create_refresh_token,
    verify_token,
    User,
    UserCreate,
    Token,
    TokenData,
    Role,
    InMemoryUserStore,
)
from shared.errors import (
    register_error_handlers,
    AuthenticationError,
    AuthorizationError,
    ValidationError,
    DatabaseError,
    ExternalAPIError,
)
from shared.secrets import get_settings

# Add parent directory to path for imports
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "../..")))

from src.core.rag_chain import DataChatRAG, RAGResponse, Message
from src.routers.query_router import QueryRouter, QueryType
from src.retrievers import DocumentRetriever
from src.cache import QueryCache

# =============================================================================
# Logging Configuration
# =============================================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# =============================================================================
# Pydantic Models
# =============================================================================


class ChatRequest(BaseModel):
    """Request model for chat endpoint."""

    question: str = Field(..., description="User's question", min_length=1, max_length=2000)
    conversation_id: Optional[str] = Field(None, description="Conversation ID for context")
    filters: Optional[Dict[str, Any]] = Field(None, description="Optional filters for document retrieval")

    model_config = {
        "json_schema_extra": {
            "examples": [
                {
                    "question": "What was our average CTR last week?",
                    "conversation_id": None,
                },
                {
                    "question": "Why is the BioGen campaign underperforming?",
                    "conversation_id": "conv_123456",
                },
            ]
        }
    }


class DocumentSource(BaseModel):
    """Document source citation."""

    content: str = Field(..., description="Snippet of the document content")
    source: str = Field(..., description="Document name/file")
    doc_type: str = Field(..., description="Document type (policy, guideline, etc.)")
    relevance: float = Field(..., description="Relevance score 0-1")


class ChatResponse(BaseModel):
    """Response model for chat endpoint."""

    answer: str = Field(..., description="Main answer text")
    query_type: str = Field(..., description="Type of query (SQL_QUERY, DOC_SEARCH, HYBRID)")
    confidence: float = Field(..., description="Classification confidence 0-1")
    conversation_id: str = Field(..., description="Conversation ID")
    sql_query: Optional[str] = Field(None, description="Generated SQL query if applicable")
    sql_results: Optional[Dict[str, Any]] = Field(None, description="SQL query results if applicable")
    doc_sources: List[DocumentSource] = Field(default_factory=list, description="Document citations")
    suggested_followup: List[str] = Field(default_factory=list, description="Suggested follow-up questions")
    processing_time_seconds: float = Field(..., description="Time to process request")
    is_cached: bool = Field(False, description="Whether the response was served from cache")
    cache_key: Optional[str] = Field(None, description="Cache key if response was cached")

    model_config = {
        "json_schema_extra": {
            "examples": [
                {
                    "answer": "Based on the data, the average CTR last week was 1.2%, which is above our healthcare benchmark of 0.8-1.5%.",
                    "query_type": "SQL_QUERY",
                    "confidence": 0.92,
                    "conversation_id": "conv_abc123",
                    "sql_query": "SELECT AVG(ctr) FROM daily_metrics WHERE date >= CURRENT_DATE - INTERVAL '7 days'",
                    "sql_results": {"avg_ctr": 1.2},
                    "doc_sources": [],
                    "suggested_followup": [
                        "💡 Would you like to see the trend over a longer time period?",
                        "💡 Would you like to compare these metrics to our benchmarks?"
                    ],
                    "processing_time_seconds": 1.23,
                }
            ]
        }
    }


class StreamEvent(BaseModel):
    """Server-Sent Event for streaming responses."""

    event: str = Field(..., description="Event type (message, error, done)")
    data: Dict[str, Any] = Field(..., description="Event data")


class ConversationMessage(BaseModel):
    """Message in conversation history."""

    role: str = Field(..., description="Message role (user, assistant)")
    content: str = Field(..., description="Message content")
    timestamp: str = Field(..., description="ISO timestamp")


class ConversationHistory(BaseModel):
    """Conversation history response."""

    conversation_id: str = Field(..., description="Conversation ID")
    messages: List[ConversationMessage] = Field(default_factory=list, description="Message history")
    created_at: str = Field(..., description="Conversation creation time")
    message_count: int = Field(..., description="Number of messages")


class DocumentUpload(BaseModel):
    """Response for document upload."""

    filename: str = Field(..., description="Uploaded filename")
    document_count: int = Field(..., description="Number of documents processed")
    chunk_count: int = Field(..., description="Number of chunks created")
    processing_time_seconds: float = Field(..., description="Time to process")
    errors: List[str] = Field(default_factory=list, description="Any errors that occurred")


class ComponentStatus(BaseModel):
    """Status of a system component."""

    name: str = Field(..., description="Component name")
    status: str = Field(..., description="Status (healthy, degraded, unhealthy)")
    message: Optional[str] = Field(None, description="Status message")
    details: Optional[Dict[str, Any]] = Field(None, description="Additional details")


class HealthResponse(BaseModel):
    """Health check response."""

    status: str = Field(..., description="Overall system status")
    version: str = Field(..., description="API version")
    timestamp: str = Field(..., description="ISO timestamp")
    components: List[ComponentStatus] = Field(default_factory=list, description="Component statuses")
    uptime_seconds: float = Field(..., description="Server uptime")


class SchemaTable(BaseModel):
    """Database table schema."""

    name: str = Field(..., description="Table name")
    description: str = Field(..., description="Table description")
    columns: List[Dict[str, str]] = Field(default_factory=list, description="Column definitions")


class SchemaResponse(BaseModel):
    """Database schema response."""

    tables: List[SchemaTable] = Field(default_factory=list, description="Available tables")
    relationships: List[Dict[str, str]] = Field(default_factory=list, description="Table relationships")


class CacheStatsResponse(BaseModel):
    """Cache statistics response."""

    cache_enabled: bool = Field(..., description="Whether caching is enabled")
    total_hits: Optional[int] = Field(None, description="Total cache hits")
    total_misses: Optional[int] = Field(None, description="Total cache misses")
    total_requests: Optional[int] = Field(None, description="Total cache requests")
    hit_rate: Optional[float] = Field(None, description="Cache hit rate percentage")
    miss_rate: Optional[float] = Field(None, description="Cache miss rate percentage")
    last_hit_at: Optional[str] = Field(None, description="Last cache hit timestamp")
    last_miss_at: Optional[str] = Field(None, description="Last cache miss timestamp")
    error: Optional[str] = Field(None, description="Error message if applicable")


class CacheClearResponse(BaseModel):
    """Cache clear response."""

    success: bool = Field(..., description="Whether the operation succeeded")
    message: str = Field(..., description="Operation message")


# =============================================================================
# In-Memory Storage (Replace with Redis in production)
# =============================================================================

class ConversationStore:
    """
    In-memory conversation storage.

    NOTE: In production, replace this with Redis or a database for:
    - Persistence across restarts
    - Distributed systems
    - TTL management
    - Better scalability
    """

    def __init__(self):
        self.conversations: Dict[str, List[Message]] = {}
        self.created_at: Dict[str, str] = {}
        self.rag_chains: Dict[str, DataChatRAG] = {}

    def get_or_create_conversation(self, conversation_id: Optional[str]) -> str:
        """Get existing conversation or create new one."""
        if conversation_id is None:
            conversation_id = f"conv_{uuid.uuid4().hex[:12]}"

        if conversation_id not in self.conversations:
            self.conversations[conversation_id] = []
            self.created_at[conversation_id] = datetime.now().isoformat()
            logger.info(f"Created new conversation: {conversation_id}")

        return conversation_id

    def add_message(self, conversation_id: str, role: str, content: str):
        """Add a message to conversation history."""
        if conversation_id in self.conversations:
            self.conversations[conversation_id].append(
                Message(role=role, content=content)
            )

    def get_messages(self, conversation_id: str) -> List[Message]:
        """Get conversation messages."""
        return self.conversations.get(conversation_id, [])

    def clear_conversation(self, conversation_id: str) -> bool:
        """Clear conversation history."""
        if conversation_id in self.conversations:
            self.conversations[conversation_id] = []
            # Also clear the RAG chain memory if it exists
            if conversation_id in self.rag_chains:
                self.rag_chains[conversation_id].clear_memory()
            return True
        return False

    def delete_conversation(self, conversation_id: str) -> bool:
        """Delete conversation entirely."""
        if conversation_id in self.conversations:
            del self.conversations[conversation_id]
            del self.created_at[conversation_id]
            if conversation_id in self.rag_chains:
                del self.rag_chains[conversation_id]
            return True
        return False

    def get_stats(self) -> Dict[str, Any]:
        """Get conversation store statistics."""
        return {
            "total_conversations": len(self.conversations),
            "total_messages": sum(len(msgs) for msgs in self.conversations.values()),
        }


# =============================================================================
# Global State
# =============================================================================

# Store for conversations
conversation_store = ConversationStore()

# RAG chain instance (will be initialized on startup)
rag_chain: Optional[DataChatRAG] = None

# API version
API_VERSION = "1.0.0"

# Server start time
start_time = time.time()


# =============================================================================
# Lifespan Management
# =============================================================================

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Manage application lifespan."""
    # Startup
    logger.info("=" * 60)
    logger.info("Starting DataChat-RAG API")
    logger.info("=" * 60)

    # Initialize RAG chain
    global rag_chain

    try:
        # Initialize components
        from src.routers import QueryRouter
        from src.retrievers import create_retriever

        # Create query router
        query_router = QueryRouter()

        # Create document retriever
        # Note: In production, use actual configuration
        doc_retriever = None
        if os.getenv("OPENAI_API_KEY"):
            doc_retriever = create_retriever(
                chroma_path=os.getenv("CHROMA_PERSIST_DIR", "./data/chromadb"),
                collection_name=os.getenv("CHROMA_COLLECTION_NAME", "datachat_documents"),
                rerank=os.getenv("RERANK_ENABLED", "false").lower() == "true",
            )
            logger.info("✓ Document retriever initialized")
        else:
            logger.warning("⚠ OPENAI_API_KEY not set, document retrieval disabled")

        # SQL retriever (to be implemented)
        sql_retriever = None
        logger.info("⚠ SQL retriever not yet implemented")

        # Initialize query cache if enabled
        query_cache = None
        enable_cache = os.getenv("CACHE_ENABLED", "true").lower() == "true"
        if enable_cache:
            try:
                redis_url = os.getenv("REDIS_URL")
                redis_host = os.getenv("REDIS_HOST", "localhost")
                redis_port = int(os.getenv("REDIS_PORT", "6379"))
                redis_db = int(os.getenv("REDIS_DB", "0"))
                redis_password = os.getenv("REDIS_PASSWORD")
                cache_ttl = int(os.getenv("CACHE_TTL", "3600"))

                if redis_url:
                    query_cache = QueryCache(
                        redis_url=redis_url,
                        default_ttl=cache_ttl,
                        enabled=True,
                    )
                else:
                    query_cache = QueryCache(
                        redis_host=redis_host,
                        redis_port=redis_port,
                        redis_db=redis_db,
                        redis_password=redis_password,
                        default_ttl=cache_ttl,
                        enabled=True,
                    )
                logger.info("✓ Query cache initialized")
            except Exception as e:
                logger.warning(f"⚠ Failed to initialize query cache: {e}")
                logger.info("  Continuing without cache (API will function normally)")

        # Create RAG chain
        rag_chain = DataChatRAG(
            sql_retriever=sql_retriever,
            doc_retriever=doc_retriever,
            query_router=query_router,
            enable_memory=True,
            query_cache=query_cache,
            enable_cache=enable_cache,
        )

        logger.info("✓ RAG chain initialized")
        logger.info("✓ API ready to accept requests")

    except Exception as e:
        logger.error(f"✗ Failed to initialize RAG chain: {e}")
        # Continue anyway for health endpoint

    yield

    # Shutdown
    logger.info("Shutting down DataChat-RAG API")
    # Close cache connection if exists
    if rag_chain and rag_chain.query_cache:
        rag_chain.query_cache.close()


# =============================================================================
# FastAPI Application
# =============================================================================

app = FastAPI(
    title="DataChat-RAG API",
    description="""
    Enterprise RAG system for healthcare AdTech internal Q&A.

    Features:
    - Query routing (SQL, documents, or hybrid)
    - Text-to-SQL for database queries
    - Semantic document retrieval
    - Conversation memory
    - Streaming responses
    - Query result caching with Redis
    """,
    version=API_VERSION,
    lifespan=lifespan,
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ORIGINS", "http://localhost:8501").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Install security filter for logs
install_security_filter()

# Register error handlers
register_error_handlers(app)

# Register rate limit exception handler
app.add_exception_handler(RateLimitExceeded, rate_limit_exception_handler)

# Get settings
settings = get_settings()


# =============================================================================
# Error Handlers
# =============================================================================

@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc):
    """Handle HTTP exceptions."""
    return JSONResponse(
        status_code=exc.status_code,
        content={
            "error": {
                "type": "http_error",
                "message": exc.detail,
                "status_code": exc.status_code,
            }
        },
    )


@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    """Handle general exceptions."""
    logger.error(f"Unhandled exception: {exc}", exc_info=True)
    return JSONResponse(
        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        content={
            "error": {
                "type": "internal_error",
                "message": "An unexpected error occurred",
                "status_code": 500,
            }
        },
    )


# =============================================================================
# Endpoints
# =============================================================================

# In-memory user store for demo (replace with database in production)
user_store = InMemoryUserStore()

# ============================================================================
# Authentication Endpoints
# ============================================================================

@app.post("/auth/register", tags=["Authentication"])
@limiter.limit("10/hour")
async def register(
    user_data: UserCreate,
    request: Request,
):
    """Register a new user."""
    try:
        user = await user_store.create_user(user_data)
        return {
            "message": "User registered successfully",
            "user": {
                "id": user.id,
                "username": user.username,
                "email": user.email,
                "role": user.role.value,
            }
        }
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )


@app.post("/auth/login", tags=["Authentication"])
@limiter.limit("20/minute")
async def login(
    username: str = Form(...),
    password: str = Form(...),
    request: Request = None,
):
    """Login and receive access token."""
    from shared.auth import authenticate_user, login_user

    user = await authenticate_user(username, password, user_store)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )

    token_data = await login_user(username, password, user_store)
    return token_data


@app.post("/auth/refresh", tags=["Authentication"])
@limiter.limit("30/minute")
async def refresh(
    refresh_token: str = Form(...),
    request: Request = None,
):
    """Refresh access token."""
    from shared.auth import refresh_user_token

    token_data = await refresh_user_token(refresh_token, user_store)
    if not token_data:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid refresh token",
            headers={"WWW-Authenticate": "Bearer"},
        )
    return token_data


@app.get("/auth/me", tags=["Authentication"])
@limiter.limit("60/minute")
async def get_current_user_info(current_user: TokenData = Depends(get_current_user)):
    """Get current user information."""
    user = await user_store.get_user_by_id(current_user.user_id)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )
    return {
        "id": user.id,
        "username": user.username,
        "email": user.email,
        "role": user.role.value,
        "is_active": user.is_active,
    }


@app.get("/", tags=["Root"])
async def root():
    """Root endpoint with API information."""
    return {
        "name": "DataChat-RAG API",
        "version": API_VERSION,
        "status": "running",
        "timestamp": datetime.now().isoformat(),
        "endpoints": {
            "chat": "/chat",
            "stream": "/chat/stream",
            "conversations": "/conversations/{conversation_id}",
            "documents": "/documents/ingest",
            "health": "/health",
            "schema": "/schema",
            "cache_stats": "/cache/stats",
            "cache_clear": "/cache/clear",
        },
    }


@app.post("/chat", response_model=ChatResponse, tags=["Chat"])
@limiter.limit("30/minute")
async def chat(request: ChatRequest, http_request: Request) -> ChatResponse:
    """
    Main chat endpoint for asking questions.

    Processes user questions through the RAG pipeline and returns
    answers with citations and metadata.
    """
    if rag_chain is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="RAG chain not initialized. Check server logs."
        )

    try:
        logger.info(f"Processing question: {request.question[:100]}...")

        # Get or create conversation
        conversation_id = conversation_store.get_or_create_conversation(request.conversation_id)

        # Process query
        response: RAGResponse = rag_chain.query(
            question=request.question,
            conversation_id=conversation_id,
            filters=request.filters,
        )

        # Update conversation history
        conversation_store.add_message(conversation_id, "user", request.question)
        conversation_store.add_message(conversation_id, "assistant", response.answer)

        # Convert to API response
        return ChatResponse(
            answer=response.answer,
            query_type=response.query_type,
            confidence=response.confidence,
            conversation_id=conversation_id,
            sql_query=response.sql_query,
            sql_results=response.sql_results.to_dict() if response.sql_results else None,
            doc_sources=[
                DocumentSource(
                    content=s["content"],
                    source=s["source"],
                    doc_type=s["doc_type"],
                    relevance=s["relevance"],
                )
                for s in response.doc_sources
            ],
            suggested_followup=response.suggested_followup,
            processing_time_seconds=response.processing_time_seconds,
            is_cached=response.is_cached,
            cache_key=response.cache_key,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing chat request: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error processing request: {str(e)}"
        )


@app.post("/chat/stream", tags=["Chat"])
async def chat_stream(request: ChatRequest):
    """
    Streaming chat endpoint using Server-Sent Events.

    Streams the response generation in real-time.
    """
    if rag_chain is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="RAG chain not initialized. Check server logs."
        )

    async def event_generator() -> AsyncGenerator[str, None]:
        """Generate SSE events."""
        try:
            conversation_id = conversation_store.get_or_create_conversation(request.conversation_id)

            # Send start event
            yield f"event: start\ndata: {json.dumps({'conversation_id': conversation_id})}\n\n"

            # Classify query
            yield "event: thinking\ndata: {\"stage\": \"classifying\"}\n\n"
            await asyncio.sleep(0.1)  # Simulate processing

            # Process query (non-streaming for now)
            response: RAGResponse = rag_chain.query(
                question=request.question,
                conversation_id=conversation_id,
                filters=request.filters,
            )

            # Send classification
            yield f"event: classified\ndata: {json.dumps({'query_type': response.query_type})}\n\n"
            await asyncio.sleep(0.1)

            # Send query type event
            if response.query_type == "SQL_QUERY":
                yield "event: thinking\ndata: {\"stage\": \"querying_database\"}\n\n"
            elif response.query_type == "DOC_SEARCH":
                yield "event: thinking\ndata: {\"stage\": \"searching_documents\"}\n\n"
            else:
                yield "event: thinking\ndata: {\"stage\": \"analyzing\"}\n\n"
            await asyncio.sleep(0.2)

            # Send chunks of the answer
            answer = response.answer
            chunk_size = 50

            for i in range(0, len(answer), chunk_size):
                chunk = answer[i:i + chunk_size]
                yield f"event: chunk\ndata: {json.dumps({'text': chunk})}\n\n"
                await asyncio.sleep(0.05)

            # Send complete event with full response
            complete_data = {
                "answer": response.answer,
                "query_type": response.query_type,
                "confidence": response.confidence,
                "sql_query": response.sql_query,
                "doc_sources": response.doc_sources,
                "suggested_followup": response.suggested_followup,
            }
            yield f"event: complete\ndata: {json.dumps(complete_data)}\n\n"

            # Update conversation history
            conversation_store.add_message(conversation_id, "user", request.question)
            conversation_store.add_message(conversation_id, "assistant", response.answer)

            # Send done event
            yield "event: done\ndata: {}\n\n"

        except Exception as e:
            logger.error(f"Error in stream: {e}", exc_info=True)
            yield f"event: error\ndata: {json.dumps({'message': str(e)})}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.get("/conversations/{conversation_id}", response_model=ConversationHistory, tags=["Conversations"])
async def get_conversation(conversation_id: str):
    """
    Get conversation history by ID.

    Returns all messages in the conversation.
    """
    messages = conversation_store.get_messages(conversation_id)

    if not messages and conversation_id not in conversation_store.created_at:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Conversation {conversation_id} not found"
        )

    return ConversationHistory(
        conversation_id=conversation_id,
        messages=[
            ConversationMessage(
                role=msg.role,
                content=msg.content,
                timestamp=msg.timestamp,
            )
            for msg in messages
        ],
        created_at=conversation_store.created_at.get(conversation_id, datetime.now().isoformat()),
        message_count=len(messages),
    )


@app.delete("/conversations/{conversation_id}", tags=["Conversations"])
async def clear_conversation(conversation_id: str):
    """
    Clear conversation history.

    Removes all messages from the conversation.
    """
    success = conversation_store.clear_conversation(conversation_id)

    if not success:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Conversation {conversation_id} not found"
        )

    return {
        "message": f"Conversation {conversation_id} cleared",
        "conversation_id": conversation_id,
    }


# Document ingest endpoint (only if python-multipart is available)
try:
    @app.post("/documents/ingest", response_model=DocumentUpload, tags=["Documents"])
    async def ingest_documents(
        background_tasks: BackgroundTasks,
        file: UploadFile = File(..., description="Document file to ingest"),
        doc_type: Optional[str] = Form(None, description="Document type (policy, guideline, etc.)"),
    ):
        """
        Upload and ingest a document into the knowledge base.

        Supports PDF, TXT, MD, DOCX, and HTML files.
        """
        if rag_chain is None or rag_chain.doc_retriever is None:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="Document retriever not initialized"
            )

        try:
            # Save uploaded file
            upload_dir = "./data/uploads"
            os.makedirs(upload_dir, exist_ok=True)

            file_path = os.path.join(upload_dir, file.filename)
            with open(file_path, "wb") as f:
                content = await file.read()
                f.write(content)

            logger.info(f"Uploaded file: {file.filename} ({len(content)} bytes)")

            # Process document in background
            def process_document():
                try:
                    from llama_index.core import Document

                    # Read file content
                    if file.filename.endswith(".pdf"):
                        # PDF processing
                        import PyPDF2
                        with open(file_path, "rb") as f:
                            reader = PyPDF2.PdfReader(f)
                            text = ""
                            for page in reader.pages:
                                text += page.extract_text() + "\n"
                    elif file.filename.endswith(".docx"):
                        import docx
                        doc = docx.Document(file_path)
                        text = "\n".join([para.text for para in doc.paragraphs])
                    else:
                        # Text-based files
                        with open(file_path, "r", encoding="utf-8") as f:
                            text = f.read()

                    # Create document with metadata
                    doc = Document(
                        text=text,
                        metadata={
                            "source": file.filename,
                            "doc_type": doc_type or "guideline",
                            "uploaded_at": datetime.now().isoformat(),
                        }
                    )

                    # Add to retriever
                    result = rag_chain.doc_retriever.add_documents([doc])

                    logger.info(f"✓ Ingested {file.filename}: {result.num_chunks} chunks")

                except Exception as e:
                    logger.error(f"Error processing document: {e}")

            background_tasks.add_task(process_document)

            return DocumentUpload(
                filename=file.filename,
                document_count=1,
                chunk_count=0,  # Will be updated after processing
                processing_time_seconds=0,
            )

        except Exception as e:
            logger.error(f"Error uploading document: {e}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error uploading document: {str(e)}"
            )

except Exception as e:
    logger.warning(f"Document ingest endpoint not available: {e}")


@app.get("/health", response_model=HealthResponse, tags=["System"])
async def health_check():
    """
    Health check endpoint.

    Returns system status and component health.
    """
    components = []

    # Check RAG chain
    if rag_chain is not None:
        components.append(ComponentStatus(
            name="rag_chain",
            status="healthy",
            message="RAG chain initialized",
        ))

        # Check document retriever
        if rag_chain.doc_retriever is not None:
            try:
                stats = rag_chain.doc_retriever.get_stats()
                components.append(ComponentStatus(
                    name="document_retriever",
                    status="healthy",
                    message=f"{stats['total_chunks']} chunks indexed",
                    details=stats,
                ))
            except Exception as e:
                components.append(ComponentStatus(
                    name="document_retriever",
                    status="unhealthy",
                    message=str(e),
                ))
        else:
            components.append(ComponentStatus(
                name="document_retriever",
                status="degraded",
                message="Not initialized (check OPENAI_API_KEY)",
            ))

        # Check SQL retriever
        if rag_chain.sql_retriever is not None:
            components.append(ComponentStatus(
                name="sql_retriever",
                status="healthy",
                message="SQL retriever initialized",
            ))
        else:
            components.append(ComponentStatus(
                name="sql_retriever",
                status="degraded",
                message="Not yet implemented",
            ))

        # Check query router
        if rag_chain.query_router is not None:
            components.append(ComponentStatus(
                name="query_router",
                status="healthy",
                message="Query router initialized",
            ))
    else:
        components.append(ComponentStatus(
            name="rag_chain",
            status="unhealthy",
            message="RAG chain not initialized",
        ))

    # Overall status
    overall_status = "healthy"
    if any(c.status == "unhealthy" for c in components):
        overall_status = "unhealthy"
    elif any(c.status == "degraded" for c in components):
        overall_status = "degraded"

    # Conversation store stats
    conv_stats = conversation_store.get_stats()
    components.append(ComponentStatus(
        name="conversation_store",
        status="healthy",
        message=f"{conv_stats['total_conversations']} conversations, {conv_stats['total_messages']} messages",
    ))

    return HealthResponse(
        status=overall_status,
        version=API_VERSION,
        timestamp=datetime.now().isoformat(),
        components=components,
        uptime_seconds=time.time() - start_time,
    )


@app.get("/schema", response_model=SchemaResponse, tags=["System"])
async def get_schema():
    """
    Get database schema information.

    Returns available tables and their columns for reference.
    """
    tables = [
        SchemaTable(
            name="campaigns",
            description="Advertising campaign metadata",
            columns=[
                {"name": "id", "type": "UUID", "description": "Unique identifier"},
                {"name": "name", "type": "VARCHAR(255)", "description": "Campaign name"},
                {"name": "client_name", "type": "VARCHAR(255)", "description": "Client company"},
                {"name": "industry", "type": "VARCHAR(50)", "description": "Industry vertical"},
                {"name": "start_date", "type": "DATE", "description": "Campaign start"},
                {"name": "end_date", "type": "DATE", "description": "Campaign end"},
                {"name": "budget", "type": "FLOAT", "description": "Budget in USD"},
                {"name": "status", "type": "VARCHAR(50)", "description": "Campaign status"},
            ],
        ),
        SchemaTable(
            name="daily_metrics",
            description="Aggregated daily campaign metrics",
            columns=[
                {"name": "date", "type": "DATE", "description": "Metric date"},
                {"name": "campaign_id", "type": "UUID", "description": "Campaign reference"},
                {"name": "impressions", "type": "INTEGER", "description": "Ad impressions"},
                {"name": "clicks", "type": "INTEGER", "description": "Click events"},
                {"name": "conversions", "type": "INTEGER", "description": "Conversion events"},
                {"name": "spend", "type": "FLOAT", "description": "Amount spent"},
                {"name": "ctr", "type": "FLOAT", "description": "Click-through rate %"},
                {"name": "cvr", "type": "FLOAT", "description": "Conversion rate %"},
                {"name": "cpa", "type": "FLOAT", "description": "Cost per acquisition"},
            ],
        ),
        SchemaTable(
            name="impressions",
            description="Individual ad impression events",
            columns=[
                {"name": "timestamp", "type": "TIMESTAMP", "description": "Event time"},
                {"name": "campaign_id", "type": "UUID", "description": "Campaign reference"},
                {"name": "device_type", "type": "VARCHAR(50)", "description": "Device category"},
                {"name": "geo_location", "type": "VARCHAR(100)", "description": "Location"},
                {"name": "ad_placement", "type": "VARCHAR(100)", "description": "Ad position"},
                {"name": "cost", "type": "FLOAT", "description": "Impression cost"},
            ],
        ),
        SchemaTable(
            name="clicks",
            description="Click events on ads",
            columns=[
                {"name": "timestamp", "type": "TIMESTAMP", "description": "Event time"},
                {"name": "impression_id", "type": "UUID", "description": "Impression reference"},
                {"name": "campaign_id", "type": "UUID", "description": "Campaign reference"},
                {"name": "landing_page", "type": "VARCHAR(500)", "description": "Destination URL"},
            ],
        ),
        SchemaTable(
            name="conversions",
            description="Conversion events (leads, signups, etc.)",
            columns=[
                {"name": "timestamp", "type": "TIMESTAMP", "description": "Event time"},
                {"name": "click_id", "type": "UUID", "description": "Click reference"},
                {"name": "campaign_id", "type": "UUID", "description": "Campaign reference"},
                {"name": "conversion_type", "type": "VARCHAR(100)", "description": "Type of conversion"},
                {"name": "value", "type": "FLOAT", "description": "Conversion value"},
            ],
        ),
    ]

    relationships = [
        {"from": "campaigns", "to": "daily_metrics", "type": "one_to_many", "key": "id=campaign_id"},
        {"from": "campaigns", "to": "impressions", "type": "one_to_many", "key": "id=campaign_id"},
        {"from": "campaigns", "to": "clicks", "type": "one_to_many", "key": "id=campaign_id"},
        {"from": "campaigns", "to": "conversions", "type": "one_to_many", "key": "id=campaign_id"},
        {"from": "impressions", "to": "clicks", "type": "one_to_many", "key": "id=impression_id"},
        {"from": "clicks", "to": "conversions", "type": "one_to_many", "key": "id=click_id"},
    ]

    return SchemaResponse(
        tables=tables,
        relationships=relationships,
    )


@app.get("/cache/stats", response_model=CacheStatsResponse, tags=["Cache"])
async def get_cache_stats():
    """
    Get cache statistics.

    Returns cache performance metrics including hit rate, miss rate, and total requests.
    """
    if rag_chain is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="RAG chain not initialized. Check server logs."
        )

    stats = rag_chain.get_cache_stats()

    return CacheStatsResponse(
        cache_enabled=stats.get("cache_enabled", False),
        total_hits=stats.get("total_hits"),
        total_misses=stats.get("total_misses"),
        total_requests=stats.get("total_requests"),
        hit_rate=stats.get("hit_rate"),
        miss_rate=stats.get("miss_rate"),
        last_hit_at=stats.get("last_hit_at"),
        last_miss_at=stats.get("last_miss_at"),
        error=stats.get("error"),
    )


@app.delete("/cache/clear", response_model=CacheClearResponse, tags=["Cache"])
async def clear_cache():
    """
    Clear the query cache.

    Removes all cached query results. Use with caution as this will impact performance
    until the cache is repopulated.
    """
    if rag_chain is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="RAG chain not initialized. Check server logs."
        )

    result = rag_chain.clear_cache()

    return CacheClearResponse(
        success=result.get("success", False),
        message=result.get("message", "Unknown result"),
    )


# =============================================================================
# Main Entry Point
# =============================================================================

def main():
    """Run the API server."""
    uvicorn.run(
        "main:app",
        host=os.getenv("API_HOST", "0.0.0.0"),
        port=int(os.getenv("API_PORT", 8000)),
        reload=os.getenv("RELOAD", "true").lower() == "true",
        workers=int(os.getenv("API_WORKERS", 1)),
    )


if __name__ == "__main__":
    import json
    main()
