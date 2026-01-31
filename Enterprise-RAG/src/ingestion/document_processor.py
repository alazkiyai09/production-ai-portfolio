# ============================================================
# Enterprise-RAG: Document Processing Pipeline
# ============================================================
"""
Intelligent document processing pipeline with multi-format support.

This module provides comprehensive document ingestion capabilities:
- Multi-format support: PDF, DOCX, MD, TXT, HTML
- Intelligent chunking with configurable size/overlap
- Metadata extraction (title, date, source, page numbers)
- Text cleaning and normalization
- Chunk ID generation for tracking

Example:
    >>> from src.ingestion import DocumentProcessor
    >>> processor = DocumentProcessor(chunk_size=512, chunk_overlap=50)
    >>> result = processor.process_file("document.pdf")
    >>> print(f"Processed {result.total_chunks} chunks")
    >>> for doc in result.documents:
    ...     print(f"Chunk {doc.chunk_id}: {doc.content[:100]}...")
"""

import hashlib
import mimetypes
import re
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from functools import wraps
from pathlib import Path
from typing import Any, BinaryIO, Optional, Tuple
from uuid import uuid4

import html2text
from docx import Document as DocxDocument
from pypdf import PdfReader

from src.config import settings
from src.exceptions import (
    DocumentChunkError,
    DocumentProcessingError,
    UnsupportedFormatError,
)
from src.logging_config import get_logger, log_execution

# Initialize logger
logger = get_logger(__name__)


# ============================================================
# Data Classes
# ============================================================

@dataclass(frozen=True)
class Document:
    """
    Represents a single document chunk with metadata.

    Attributes:
        content: The text content of the chunk
        metadata: Associated metadata (source, page, chunk index, etc.)
        doc_id: Unique identifier for the source document
        chunk_id: Unique identifier for this specific chunk

    Example:
        >>> doc = Document(
        ...     content="This is the content",
        ...     metadata={"source": "file.pdf", "page": 1},
        ...     doc_id="doc_123",
        ...     chunk_id="doc_123_chunk_0"
        ... )
    """

    content: str
    metadata: dict[str, Any]
    doc_id: str
    chunk_id: str

    def __post_init__(self) -> None:
        """Validate document after initialization."""
        if not self.content or not self.content.strip():
            raise ValueError("Document content cannot be empty")
        if not self.doc_id:
            raise ValueError("Document ID cannot be empty")
        if not self.chunk_id:
            raise ValueError("Chunk ID cannot be empty")

    def to_dict(self) -> dict[str, Any]:
        """
        Convert document to dictionary format.

        Returns:
            Dictionary representation of the document

        Example:
            >>> doc_dict = doc.to_dict()
            >>> print(doc_dict['content'])
        """
        return {
            "content": self.content,
            "metadata": self.metadata,
            "doc_id": self.doc_id,
            "chunk_id": self.chunk_id,
        }

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "Document":
        """
        Create Document from dictionary.

        Args:
            data: Dictionary with document data

        Returns:
            Document instance

        Example:
            >>> doc = Document.from_dict(doc_dict)
        """
        return cls(
            content=data["content"],
            metadata=data["metadata"],
            doc_id=data["doc_id"],
            chunk_id=data["chunk_id"],
        )


@dataclass
class ProcessingResult:
    """
    Result of document processing operation.

    Attributes:
        documents: List of processed document chunks
        total_chunks: Total number of chunks created
        processing_time: Time taken to process in seconds
        errors: List of error messages (if any)
        files_processed: Number of files processed
        files_failed: Number of files that failed to process

    Example:
        >>> result = processor.process_file("doc.pdf")
        >>> print(f"Processed {result.total_chunks} chunks in {result.processing_time:.2f}s")
        >>> if result.errors:
        ...     for error in result.errors:
        ...         print(f"Error: {error}")
    """

    documents: list[Document] = field(default_factory=list)
    total_chunks: int = 0
    processing_time: float = 0.0
    errors: list[str] = field(default_factory=list)
    files_processed: int = 0
    files_failed: int = 0

    def to_dict(self) -> dict[str, Any]:
        """Convert processing result to dictionary."""
        return {
            "documents": [doc.to_dict() for doc in self.documents],
            "total_chunks": self.total_chunks,
            "processing_time": self.processing_time,
            "errors": self.errors,
            "files_processed": self.files_processed,
            "files_failed": self.files_failed,
        }


# ============================================================
# Supported File Types
# ============================================================

class FileType(str, Enum):
    """Supported document file types."""

    PDF = "pdf"
    DOCX = "docx"
    MARKDOWN = "md"
    TEXT = "txt"
    HTML = "html"


# ============================================================
# Document Processor
# ============================================================

class DocumentProcessor:
    """
    Intelligent document processing pipeline.

    Handles multi-format document ingestion with intelligent chunking,
    metadata extraction, and text normalization.

    Features:
        - Multi-format support (PDF, DOCX, MD, TXT, HTML)
        - Configurable chunk size and overlap
        - Metadata extraction
        - Text cleaning and normalization
        - Paragraph-aware chunking
        - Header context preservation

    Example:
        >>> processor = DocumentProcessor(chunk_size=512, chunk_overlap=50)
        >>> result = processor.process_file("document.pdf")
        >>> for doc in result.documents:
        ...     print(f"{doc.chunk_id}: {doc.content[:50]}...")
    """

    # File type mappings
    MIME_TYPE_MAP = {
        "application/pdf": FileType.PDF,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": FileType.DOCX,
        "text/markdown": FileType.MARKDOWN,
        "text/plain": FileType.TEXT,
        "text/html": FileType.HTML,
    }

    EXTENSION_MAP = {
        ".pdf": FileType.PDF,
        ".docx": FileType.DOCX,
        ".doc": FileType.DOCX,
        ".md": FileType.MARKDOWN,
        ".markdown": FileType.MARKDOWN,
        ".txt": FileType.TEXT,
        ".html": FileType.HTML,
        ".htm": FileType.HTML,
    }

    def __init__(
        self,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
        min_chunk_size: int = 100,
        preserve_headers: bool = True,
        header_context_size: int = 100,
    ) -> None:
        """
        Initialize the document processor.

        Args:
            chunk_size: Target size for each chunk in characters
            chunk_overlap: Overlap between consecutive chunks
            min_chunk_size: Minimum size for a valid chunk
            preserve_headers: Whether to add document header to each chunk
            header_context_size: Number of characters from doc start for context
        """
        self.chunk_size = chunk_size or settings.CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
        self.min_chunk_size = min_chunk_size
        self.preserve_headers = preserve_headers
        self.header_context_size = header_context_size

        # Validate configuration
        if self.chunk_size <= self.chunk_overlap:
            raise ValueError(
                f"chunk_size ({self.chunk_size}) must be greater than "
                f"chunk_overlap ({self.chunk_overlap})"
            )
        if self.min_chunk_size >= self.chunk_size:
            raise ValueError(
                f"min_chunk_size ({self.min_chunk_size}) must be less than "
                f"chunk_size ({self.chunk_size})"
            )

        # Initialize HTML to text converter
        self.html_converter = html2text.HTML2Text()
        self.html_converter.ignore_links = False
        self.html_converter.ignore_images = False
        self.html_converter.body_width = 0  # Don't wrap lines

        logger.info(
            "DocumentProcessor initialized",
            extra={
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap,
                "min_chunk_size": self.min_chunk_size,
            },
        )

    @log_execution(include_args=False)
    def process_file(self, file_path: Path | str) -> ProcessingResult:
        """
        Process a single file into document chunks.

        Args:
            file_path: Path to the file to process

        Returns:
            ProcessingResult with document chunks

        Raises:
            DocumentProcessingError: If file cannot be processed
            UnsupportedFormatError: If file format is not supported

        Example:
            >>> result = processor.process_file("document.pdf")
            >>> print(f"Created {result.total_chunks} chunks")
        """
        import time

        start_time = time.time()
        file_path = Path(file_path)

        # Validate file exists
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Detect file type
        file_type = self._detect_file_type(file_path)

        logger.info(
            f"Processing file: {file_path.name}",
            extra={"file_path": str(file_path), "file_type": file_type},
        )

        try:
            # Extract text and metadata based on file type
            if file_type == FileType.PDF:
                text, metadata = self._extract_text_pdf(file_path)
            elif file_type == FileType.DOCX:
                text, metadata = self._extract_text_docx(file_path)
            elif file_type == FileType.MARKDOWN:
                text, metadata = self._extract_text_markdown(file_path)
            elif file_type == FileType.HTML:
                text, metadata = self._extract_text_html(file_path)
            elif file_type == FileType.TEXT:
                text, metadata = self._extract_text_text(file_path)
            else:
                raise UnsupportedFormatError(
                    file_type=file_type,
                    supported_formats=[ft.value for ft in FileType],
                )

            # Clean the extracted text
            text = self._clean_text(text)

            # Validate we have content
            if not text or len(text.strip()) < 10:
                raise DocumentProcessingError(
                    message=f"Insufficient text content extracted from {file_path.name}",
                    file_path=str(file_path),
                )

            # Generate document ID
            doc_id = self._generate_doc_id(file_path)

            # Add common metadata
            metadata.update({
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": file_type,
                "doc_id": doc_id,
                "processed_at": datetime.utcnow().isoformat(),
            })

            # Get header context if enabled
            header_context = ""
            if self.preserve_headers:
                header_context = text[: self.header_context_size]
                metadata["header_context"] = header_context

            # Chunk the text
            documents = self._chunk_text(text, metadata, doc_id)

            processing_time = time.time() - start_time

            result = ProcessingResult(
                documents=documents,
                total_chunks=len(documents),
                processing_time=processing_time,
                errors=[],
                files_processed=1,
                files_failed=0,
            )

            logger.info(
                f"Successfully processed {file_path.name}",
                extra={
                    "chunks": len(documents),
                    "processing_time": round(processing_time, 2),
                },
            )

            return result

        except Exception as e:
            processing_time = time.time() - start_time
            error_msg = f"Failed to process {file_path.name}: {str(e)}"

            logger.error(error_msg, exc_info=True)

            return ProcessingResult(
                documents=[],
                total_chunks=0,
                processing_time=processing_time,
                errors=[error_msg],
                files_processed=0,
                files_failed=1,
            )

    @log_execution(include_args=False)
    def process_directory(
        self,
        dir_path: Path | str,
        recursive: bool = True,
        pattern: str = "*",
    ) -> ProcessingResult:
        """
        Process all supported files in a directory.

        Args:
            dir_path: Path to the directory
            recursive: Whether to process subdirectories
            pattern: Glob pattern for file matching

        Returns:
            ProcessingResult with all document chunks

        Example:
            >>> result = processor.process_directory("./documents", recursive=True)
            >>> print(f"Processed {result.files_processed} files")
        """
        import time

        start_time = time.time()
        dir_path = Path(dir_path)

        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {dir_path}")

        # Find all files
        if recursive:
            files = list(dir_path.rglob(pattern))
        else:
            files = list(dir_path.glob(pattern))

        # Filter to supported files
        supported_extensions = set(self.EXTENSION_MAP.keys())
        files = [f for f in files if f.is_file() and f.suffix.lower() in supported_extensions]

        logger.info(
            f"Processing directory: {dir_path}",
            extra={"file_count": len(files), "recursive": recursive},
        )

        # Process all files
        all_documents: list[Document] = []
        all_errors: list[str] = []
        files_processed = 0
        files_failed = 0

        for file_path in files:
            try:
                result = self.process_file(file_path)
                all_documents.extend(result.documents)
                all_errors.extend(result.errors)
                files_processed += result.files_processed
                files_failed += result.files_failed
            except Exception as e:
                error_msg = f"Failed to process {file_path.name}: {str(e)}"
                all_errors.append(error_msg)
                files_failed += 1
                logger.error(error_msg, exc_info=True)

        processing_time = time.time() - start_time

        logger.info(
            f"Directory processing complete",
            extra={
                "files_processed": files_processed,
                "files_failed": files_failed,
                "total_chunks": len(all_documents),
                "processing_time": round(processing_time, 2),
            },
        )

        return ProcessingResult(
            documents=all_documents,
            total_chunks=len(all_documents),
            processing_time=processing_time,
            errors=all_errors,
            files_processed=files_processed,
            files_failed=files_failed,
        )

    @log_execution(include_args=False)
    def process_bytes(
        self,
        content: bytes,
        filename: str,
        metadata: Optional[dict[str, Any]] = None,
    ) -> ProcessingResult:
        """
        Process file from bytes (for API uploads).

        Args:
            content: File content as bytes
            filename: Original filename
            metadata: Additional metadata to include

        Returns:
            ProcessingResult with document chunks

        Example:
            >>> with open("document.pdf", "rb") as f:
            ...     content = f.read()
            >>> result = processor.process_bytes(content, "document.pdf")
        """
        import tempfile
        import time

        start_time = time.time()

        # Detect file type from filename
        file_path = Path(filename)
        file_type = self._detect_file_type(file_path)

        logger.info(
            f"Processing bytes: {filename}",
            extra={"filename": filename, "file_type": file_type, "size": len(content)},
        )

        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(
                delete=False,
                suffix=file_path.suffix,
            ) as temp_file:
                temp_file.write(content)
                temp_path = Path(temp_file.name)

            # Process the temporary file
            result = self.process_file(temp_path)

            # Update metadata with original filename and any provided metadata
            for doc in result.documents:
                doc.metadata["filename"] = filename
                if metadata:
                    doc.metadata.update(metadata)

            # Clean up temp file
            temp_path.unlink(missing_ok=True)

            return result

        except Exception as e:
            logger.error(f"Failed to process bytes for {filename}: {str(e)}", exc_info=True)

            return ProcessingResult(
                documents=[],
                total_chunks=0,
                processing_time=time.time() - start_time,
                errors=[f"Failed to process {filename}: {str(e)}"],
                files_processed=0,
                files_failed=1,
            )

    # ============================================================
    # File Type Detection
    # ============================================================

    def _detect_file_type(self, file_path: Path) -> FileType:
        """
        Detect file type from extension and MIME type.

        Args:
            file_path: Path to the file

        Returns:
            FileType enum value

        Raises:
            UnsupportedFormatError: If file type is not supported
        """
        # Try extension first
        ext = file_path.suffix.lower()
        if ext in self.EXTENSION_MAP:
            return self.EXTENSION_MAP[ext]

        # Try MIME type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if mime_type in self.MIME_TYPE_MAP:
            return self.MIME_TYPE_MAP[mime_type]

        # Not supported
        raise UnsupportedFormatError(
            file_type=ext or "unknown",
            supported_formats=list(self.EXTENSION_MAP.keys()),
        )

    # ============================================================
    # Text Extraction Methods
    # ============================================================

    def _extract_text_pdf(self, file_path: Path) -> Tuple[str, dict[str, Any]]:
        """
        Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple of (extracted_text, metadata)

        Example:
            >>> text, metadata = processor._extract_text_pdf(Path("doc.pdf"))
        """
        try:
            reader = PdfReader(str(file_path))
            metadata = {
                "pages": len(reader.pages),
                "file_type": FileType.PDF,
            }

            # Extract PDF metadata if available
            if reader.metadata:
                if reader.metadata.get("/Title"):
                    metadata["title"] = reader.metadata.get("/Title")
                if reader.metadata.get("/Author"):
                    metadata["author"] = reader.metadata.get("/Author")
                if reader.metadata.get("/CreationDate"):
                    metadata["creation_date"] = reader.metadata.get("/CreationDate")

            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages, start=1):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
                except Exception as e:
                    logger.warning(
                        f"Failed to extract page {page_num} from {file_path.name}: {e}"
                    )

            text = "\n\n".join(text_parts)

            logger.debug(
                f"Extracted {len(text)} characters from PDF with {metadata['pages']} pages"
            )

            return text, metadata

        except Exception as e:
            raise DocumentProcessingError(
                message=f"Failed to extract text from PDF: {str(e)}",
                file_path=str(file_path),
                file_type="pdf",
            )

    def _extract_text_docx(self, file_path: Path) -> Tuple[str, dict[str, Any]]:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            doc = DocxDocument(str(file_path))

            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "file_type": FileType.DOCX,
            }

            # Extract core properties
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            if doc.core_properties.created:
                metadata["creation_date"] = doc.core_properties.created.isoformat()

            # Extract paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    text_parts.append(para.text)

            # Extract tables (add as structured text)
            for table_num, table in enumerate(doc.tables, start=1):
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_parts.append(f"\n[Table {table_num}]\n" + "\n".join(table_text))

            text = "\n\n".join(text_parts)

            logger.debug(
                f"Extracted {len(text)} characters from DOCX with "
                f"{metadata['paragraphs']} paragraphs and {metadata['tables']} tables"
            )

            return text, metadata

        except Exception as e:
            raise DocumentProcessingError(
                message=f"Failed to extract text from DOCX: {str(e)}",
                file_path=str(file_path),
                file_type="docx",
            )

    def _extract_text_markdown(self, file_path: Path) -> Tuple[str, dict[str, Any]]:
        """
        Extract text from Markdown file.

        Args:
            file_path: Path to Markdown file

        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            # Read with encoding detection
            text = self._read_file_with_encoding(file_path)

            # Extract title (first # heading)
            title = None
            for line in text.split("\n"):
                if line.strip().startswith("# "):
                    title = line.strip().replace("# ", "", 1)
                    break

            # Count headings
            heading_count = len(re.findall(r"^#{1,6}\s", text, re.MULTILINE))

            metadata = {
                "file_type": FileType.MARKDOWN,
                "headings": heading_count,
            }
            if title:
                metadata["title"] = title

            logger.debug(
                f"Extracted {len(text)} characters from Markdown with {heading_count} headings"
            )

            return text, metadata

        except Exception as e:
            raise DocumentProcessingError(
                message=f"Failed to extract text from Markdown: {str(e)}",
                file_path=str(file_path),
                file_type="md",
            )

    def _extract_text_html(self, file_path: Path) -> Tuple[str, dict[str, Any]]:
        """
        Extract text from HTML file.

        Args:
            file_path: Path to HTML file

        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            html_content = self._read_file_with_encoding(file_path)

            # Convert HTML to markdown/text
            text = self.html_converter.handle(html_content)

            # Extract title from HTML if present
            title = None
            title_match = re.search(r"<title>(.*?)</title>", html_content, re.IGNORECASE)
            if title_match:
                title = title_match.group(1).strip()

            metadata = {
                "file_type": FileType.HTML,
            }
            if title:
                metadata["title"] = title

            logger.debug(f"Extracted {len(text)} characters from HTML")

            return text, metadata

        except Exception as e:
            raise DocumentProcessingError(
                message=f"Failed to extract text from HTML: {str(e)}",
                file_path=str(file_path),
                file_type="html",
            )

    def _extract_text_text(self, file_path: Path) -> Tuple[str, dict[str, Any]]:
        """
        Extract text from plain text file.

        Args:
            file_path: Path to text file

        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            text = self._read_file_with_encoding(file_path)

            metadata = {
                "file_type": FileType.TEXT,
                "characters": len(text),
                "lines": len(text.split("\n")),
            }

            logger.debug(f"Extracted {len(text)} characters from text file")

            return text, metadata

        except Exception as e:
            raise DocumentProcessingError(
                message=f"Failed to extract text from file: {str(e)}",
                file_path=str(file_path),
                file_type="txt",
            )

    def _read_file_with_encoding(self, file_path: Path) -> str:
        """
        Read file with automatic encoding detection.

        Args:
            file_path: Path to file

        Returns:
            File content as string
        """
        # Try common encodings
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"]

        for encoding in encodings:
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue

        # If all fail, try with errors='replace'
        return file_path.read_text(encoding="utf-8", errors="replace")

    # ============================================================
    # Text Cleaning
    # ============================================================

    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize text.

        Operations:
        - Remove excessive whitespace
        - Fix common encoding issues
        - Remove control characters
        - Normalize line breaks
        - Remove zero-width characters

        Args:
            text: Raw text to clean

        Returns:
            Cleaned text

        Example:
            >>> clean = processor._clean_text("  Extra   spaces  ")
            >>> assert clean == "Extra spaces"
        """
        if not text:
            return ""

        # Remove zero-width characters
        text = re.sub(r"[\u200b\u200c\u200d\u2060\ufeff]", "", text)

        # Remove control characters except newlines and tabs
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

        # Normalize line breaks
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"\r", "\n", text)

        # Remove excessive blank lines (more than 2 consecutive)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Remove excessive spaces within lines
        text = re.sub(r" {3,}", "  ", text)

        # Fix common encoding issues
        text = text.replace("\x92", "'")  # Smart quote
        text = text.replace("\x93", '"')
        text = text.replace("\x94", '"')
        text = text.replace("\x85", "...")  # Ellipsis

        # Strip leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)

        # Clean up final result
        text = text.strip()

        return text

    # ============================================================
    # Text Chunking
    # ============================================================

    def _chunk_text(
        self,
        text: str,
        metadata: dict[str, Any],
        doc_id: str,
    ) -> list[Document]:
        """
        Split text into overlapping chunks preserving paragraph boundaries.

        Strategy:
        1. Try to split on paragraph boundaries first
        2. If a paragraph is too long, split on sentence boundaries
        3. If a sentence is too long, split on word boundaries
        4. Add overlap between chunks for context

        Args:
            text: Text to chunk
            metadata: Metadata to attach to each chunk
            doc_id: Document ID

        Returns:
            List of Document chunks

        Raises:
            DocumentChunkError: If chunking fails
        """
        try:
            chunks: list[Document] = []

            # Split into paragraphs
            paragraphs = text.split("\n\n")

            current_chunk = ""
            chunk_index = 0
            char_count = 0

            for para_idx, paragraph in enumerate(paragraphs):
                paragraph = paragraph.strip()
                if not paragraph:
                    continue

                para_len = len(paragraph)

                # If adding this paragraph would exceed chunk size
                if char_count + para_len > self.chunk_size and current_chunk:
                    # Save current chunk
                    chunk_text = current_chunk.strip()
                    if len(chunk_text) >= self.min_chunk_size:
                        chunk = self._create_chunk(
                            chunk_text,
                            metadata,
                            doc_id,
                            chunk_index,
                        )
                        chunks.append(chunk)
                        chunk_index += 1

                    # Start new chunk with overlap
                    overlap_text = self._get_overlap_text(current_chunk)
                    current_chunk = overlap_text + "\n\n" + paragraph
                    char_count = len(current_chunk)
                else:
                    # Add paragraph to current chunk
                    if current_chunk:
                        current_chunk += "\n\n" + paragraph
                    else:
                        current_chunk = paragraph
                    char_count += para_len + 2  # +2 for "\n\n"

                # Handle very long paragraphs
                while char_count > self.chunk_size:
                    # Split the current chunk
                    split_point = self.chunk_size - self.chunk_overlap

                    # Find good break point
                    break_point = self._find_break_point(current_chunk, split_point)

                    # Create chunk from first part
                    chunk_text = current_chunk[:break_point].strip()
                    if len(chunk_text) >= self.min_chunk_size:
                        chunk = self._create_chunk(
                            chunk_text,
                            metadata,
                            doc_id,
                            chunk_index,
                        )
                        chunks.append(chunk)
                        chunk_index += 1

                    # Keep second part for next chunk
                    current_chunk = current_chunk[break_point:].strip()
                    char_count = len(current_chunk)

            # Add final chunk
            if current_chunk.strip() and len(current_chunk.strip()) >= self.min_chunk_size:
                chunk = self._create_chunk(
                    current_chunk.strip(),
                    metadata,
                    doc_id,
                    chunk_index,
                )
                chunks.append(chunk)

            # Add chunk metadata
            for idx, chunk in enumerate(chunks):
                chunk.metadata["chunk_index"] = idx
                chunk.metadata["total_chunks"] = len(chunks)

            logger.debug(f"Created {len(chunks)} chunks from document {doc_id}")

            return chunks

        except Exception as e:
            raise DocumentChunkError(
                message=f"Failed to chunk text: {str(e)}",
                chunk_size=self.chunk_size,
            )

    def _create_chunk(
        self,
        content: str,
        metadata: dict[str, Any],
        doc_id: str,
        chunk_index: int,
    ) -> Document:
        """
        Create a Document chunk with metadata.

        Args:
            content: Chunk content
            metadata: Base metadata
            doc_id: Document ID
            chunk_index: Chunk index

        Returns:
            Document instance
        """
        chunk_id = self._generate_chunk_id(doc_id, chunk_index)

        # Create chunk-specific metadata
        chunk_metadata = metadata.copy()
        chunk_metadata.update({
            "chunk_index": chunk_index,
            "chunk_id": chunk_id,
            "char_count": len(content),
        })

        # Add header context if available and not first chunk
        if (
            chunk_index > 0
            and "header_context" in metadata
            and metadata["header_context"]
        ):
            chunk_metadata["has_header_context"] = True

        return Document(
            content=content,
            metadata=chunk_metadata,
            doc_id=doc_id,
            chunk_id=chunk_id,
        )

    def _get_overlap_text(self, text: str) -> str:
        """
        Get overlap text from end of current chunk.

        Args:
            text: Current chunk text

        Returns:
            Overlap portion of text
        """
        if len(text) <= self.chunk_overlap:
            return text

        # Find a good break point
        break_point = self._find_break_point(text, len(text) - self.chunk_overlap)

        return text[break_point:]

    def _find_break_point(self, text: str, target_pos: int) -> int:
        """
        Find a good break point near target position.

        Prefers:
        1. Sentence boundaries (. ! ?)
        2. Word boundaries (spaces)
        3. Falls back to exact position

        Args:
            text: Text to search
            target_pos: Target position

        Returns:
            Actual break position
        """
        # Search range around target
        search_range = 100  # Look 100 chars forward/backward
        start = max(0, target_pos - search_range)
        end = min(len(text), target_pos + search_range)
        search_text = text[start:end]

        # Try sentence boundaries first
        for delimiter in (". ", "! ", "? "):
            # Look forward from target
            rel_pos = target_pos - start
            pos = search_text.find(delimiter, rel_pos)
            if pos != -1 and abs(pos - rel_pos) < search_range:
                return start + pos + len(delimiter)

            # Look backward from target
            pos = search_text.rfind(delimiter, 0, rel_pos)
            if pos != -1 and abs(pos - rel_pos) < search_range:
                return start + pos + len(delimiter)

        # Try word boundaries
        for delimiter in (" ", "\n"):
            rel_pos = target_pos - start
            pos = search_text.find(delimiter, rel_pos)
            if pos != -1 and abs(pos - rel_pos) < 20:
                return start + pos + len(delimiter)

            pos = search_text.rfind(delimiter, 0, rel_pos)
            if pos != -1 and abs(pos - rel_pos) < 20:
                return start + pos + len(delimiter)

        # Fall back to target position
        return target_pos

    # ============================================================
    # ID Generation
    # ============================================================

    def _generate_doc_id(self, file_path: Path) -> str:
        """
        Generate unique document ID from file path.

        Args:
            file_path: Path to file

        Returns:
            Unique document ID

        Example:
            >>> doc_id = processor._generate_doc_id(Path("document.pdf"))
            >>> assert doc_id.startswith("doc_")
        """
        # Create hash from file path and size
        file_stat = file_path.stat()
        hash_input = f"{file_path}{file_stat.st_size}{file_stat.st_mtime}"

        hash_obj = hashlib.md5(hash_input.encode())
        return f"doc_{hash_obj.hexdigest()[:12]}"

    def _generate_chunk_id(self, doc_id: str, chunk_index: int) -> str:
        """
        Generate unique chunk ID.

        Args:
            doc_id: Document ID
            chunk_index: Chunk index

        Returns:
            Unique chunk ID

        Example:
            >>> chunk_id = processor._generate_chunk_id("doc_abc123", 0)
            >>> assert chunk_id == "doc_abc123_chunk_0"
        """
        return f"{doc_id}_chunk_{chunk_index}"


# ============================================================
# Utility Functions
# ============================================================

def create_processor_from_settings() -> DocumentProcessor:
    """
    Create DocumentProcessor configured from settings.

    Returns:
        Configured DocumentProcessor instance

    Example:
        >>> processor = create_processor_from_settings()
        >>> result = processor.process_file("doc.pdf")
    """
    return DocumentProcessor(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
    )


# Export public API
__all__ = [
    # Data classes
    "Document",
    "ProcessingResult",
    # Main class
    "DocumentProcessor",
    # Enums
    "FileType",
    # Utilities
    "create_processor_from_settings",
]
